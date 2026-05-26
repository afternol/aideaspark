# AideaSpark アイデア監査方法をWordファイルに出力するスクリプト
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"C:\Users\after\AI副業_YouTube"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "AideaSpark_アイデア監査手法.docx")

doc = Document()

# ページ余白設定
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# スタイル定義
def heading1(text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def note(text):
    """グレー背景の注記ボックス"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # ヘッダー行の背景色
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "DBEAFE")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def pagebreak():
    doc.add_page_break()

# =====================================================================
# タイトル
# =====================================================================
title = doc.add_heading("AideaSpark アイデア監査（ファクトチェック）手法ガイド", level=0)
title.runs[0].font.size = Pt(20)
title.runs[0].font.color.rgb = RGBColor(0x11, 0x3D, 0xA0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("作成日：2026年4月28日　　対象データ：全200件ビジネスアイデア（idea-1〜200）")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# =====================================================================
# 1. 概要
# =====================================================================
heading1("1. 監査の概要と実施成果")

body("AideaSparkに登録された200件のビジネスアイデアに対して、複数ラウンドの体系的ファクトチェックを実施した。"
     "累計233件以上の誤りを発見・修正し、データ品質を大幅に向上させた。")

doc.add_paragraph()
add_table(
    ["区分", "対象", "修正件数", "手法"],
    [
        ["R1〜R11（旧100件）", "idea-1〜100", "約50件以上", "10並列エージェント × WebSearch × Supabase直接UPDATE"],
        ["第2〜8ラウンド", "idea-101〜200", "約170件", "10並列エージェント × WebSearch × Supabase直接UPDATE"],
        ["インサイトパネル専用", "全200件（3フィールド）", "18件", "10並列エージェント × WebSearch（1件2回以上）"],
        ["strengthNote専用", "全200件", "0件直接+4件隣接", "10並列エージェント × WebSearch"],
        ["累計", "全200件", "233件以上", "—"],
    ],
    [3.5, 5.0, 2.5, 7.5]
)

# =====================================================================
# 2. 監査プロセス
# =====================================================================
pagebreak()
heading1("2. 標準監査プロセス")

heading2("2-1. グループ分割と並列実行")
body("100件を10グループ（各10件）に分割し、10エージェントを同時起動する。"
     "1ラウンドで全100件を網羅でき、所要時間は約5〜10分。")

heading2("2-2. エージェントへの指示テンプレート（重要）")
body("各エージェントに以下の検証観点を必ず指示する：", bold=True)
bullet("出典不明の具体的数値（%・万人・億円）→ WebSearchで一次ソースまで遡って確認")
bullet("根拠なき断言の実態確認（「急増」「唯一」「最大」等）")
bullet("競合サービスの機能・料金・現状（廃止・統合・料金変更）")
bullet("法規制・義務化の正確な施行状況（法律名・条文・施行日を確認）")
bullet("企業名・サービス名の正式名称（旧名・略称の誤用）")
bullet("WebSearchで裏付けが取れない主張は削除・修正")

doc.add_paragraph()
heading2("2-3. DB直接UPDATE構文")
body("Supabase（PostgreSQL）への直接UPDATEスクリプト例：", bold=True)
code = doc.add_paragraph()
code.paragraph_format.left_indent = Cm(1.0)
run = code.add_run(
    "import pg from 'pg';\n"
    "const pool = new pg.Pool({ connectionString: 'postgresql://...' });\n"
    "await pool.query(\n"
    "  'UPDATE \"Idea\" SET \"フィールド名\" = $1, \"updatedAt\" = NOW() WHERE number = $2',\n"
    "  [新しい値, 番号]\n"
    ");\n"
    "await pool.end();"
)
run.font.size = Pt(9)
run.font.name = "Courier New"
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

heading2("2-4. 報告フォーマット")
body("各エージェントは以下のフォーマットで報告する：")
bullet("修正あり：#番号 フィールド名: 旧値→新値（理由・出典URL）")
bullet("修正なし：#番号: 確認済み（検証した主要ポイントと出典）")

heading2("2-5. フィールド別チェックの優先順位")
add_table(
    ["フィールド", "リスク", "主な問題パターン"],
    [
        ["competitors", "★★★★★", "廃止サービス混入・料金陳腐化・企業名誤り"],
        ["whyNow", "★★★★★", "義務化誤用・古い統計・施行日誤り"],
        ["noveltyNote", "★★★★★", "廃止サービスを「存在しない」の根拠に使用"],
        ["problem", "★★★★☆", "数値の根拠不明・対象範囲誤り"],
        ["scoreComments", "★★★★☆", "市場規模の誇張・数値の時点誤記"],
        ["competitiveEdge", "★★★☆☆", "競合機能の過小評価・「唯一」の誤用"],
        ["strengthNote", "★☆☆☆☆", "抽象的優位性記述のため誤り率が構造的に低い"],
    ],
    [4.0, 2.0, 12.0]
)

# =====================================================================
# 3. ハルシネーションパターン分類
# =====================================================================
pagebreak()
heading1("3. 発見されたハルシネーションパターン（A〜ψ）")

body("計20以上の類型を体系化。新しい監査ラウンドで追加パターンを発見次第、随時更新する。")
doc.add_paragraph()

heading2("【カテゴリ1】サービス・競合情報の誤り")

heading3("A. 廃止・終了済みサービスの混入（最多・最重大）")
body("競合として記載されているサービスが実際にはすでに終了している。")
add_table(
    ["誤記例", "正確な情報"],
    [
        ["SEND（競合として記載）", "2022年破産"],
        ["Quotta（競合として記載）", "2025年終了"],
        ["LINEヘルスケア（競合として記載）", "2025年6月終了"],
        ["FiNC（競合として記載）", "2026年9月終了予定"],
        ["ぐるなび仕入モール（別名で参照）", "一般消費者向け「ぐるなび食市場」と混同"],
        ["ペットみるん（noveltyNoteの根拠）", "2020年5月31日サービス終了"],
    ],
    [7.0, 11.5]
)
note("対策：competitors フィールドは定期的に稼働確認（現在の公式URL・プレスリリース）と買収履歴の検索を実施する。")

doc.add_paragraph()
heading3("E. 企業名・サービス名の誤記（リブランド・統合未反映）")
add_table(
    ["誤記", "正確な名称・理由"],
    [
        ["HiPro Biz", "HiPro Direct（HiPro Bizは経営幹部向け別サービス）"],
        ["GVA assist", "OLGA（2024年11月ブランド統合）"],
        ["SPEEDA・FORCAS", "「スピーダ」に3ブランド統一（2024年7月）"],
        ["LegalForce", "LegalOn（製品名LegalForceは現在も継続、会社名がLegalOn Technologiesに変更）"],
        ["フリーナンス（GMO）", "FREENANCE by freee（2025年10月にGMOからfreeeへ統合）"],
        ["VoiceText翻訳", "VoiceTextは音声合成エンジン（翻訳SaaSではない）"],
        ["BAONZ", "BATONZ（事業承継M&AプラットフォームのB正式表記）"],
        ["Nuance DAX", "Microsoft Dragon Copilot（2025年3月リブランド）"],
        ["PayPay資産運用", "PayPay証券（2026年2月9日名称変更）"],
    ],
    [5.5, 13.0]
)
note("対策：M&A・リブランド・社名変更は頻繁。最新プレスリリースを確認する。")

doc.add_paragraph()
heading3("G. 競合サービスの機能・料金の陳腐化")
body("競合説明が古い料金・機能情報に固定されているパターン。")
add_table(
    ["競合名", "誤った記述", "正確な情報"],
    [
        ["Design Pickle", "旧ティア制", "2025年6月に新モデル（Base $119+Creative Hours）へ移行"],
        ["Penji", "「$99〜」", "「$499〜$1,497/月」の3段階に変更"],
        ["ChatGPT Deep Research", "「Plusプランで無料」", "月額20ドルのPlusプランの有料機能・月25回まで"],
        ["Mazrica", "機能未記載", "2025年11月Mazrica Target（550万社DB＋提案書1クリック生成）リリース"],
    ],
    [4.0, 5.5, 9.0]
)
note("対策：価格・機能情報は最低6ヶ月ごとに再確認。競合の公式ブログ・PR TIMESで追跡。")

doc.add_paragraph()
heading3("ι. 架空サービス名の競合欄混入")
body("実在しないサービス名が競合として記載されている。")
bullet("「ディスカバー農業」: 実在しない農業マッチングサービス")
bullet("「横浜生命保険」: 実在しない保険会社")
bullet("「みんなの法人保険」: 実在確認不能→削除")
bullet("「ビアジャパン」「MarketBoard」: 実在確認不能→実在する競合に差替え")
note("対策：固有名詞（企業・サービス名）は必ずWebSearchで公式サイトの存在を確認する。")

# =====================================================================
heading2("【カテゴリ2】法令・規制の誤り")

heading3("B. 「義務化」「必須」の誤用")
body("任意・推奨・ガイドラインを「義務化」と断言するパターン。")
add_table(
    ["誤記例", "正確な情報"],
    [
        ["Scope3開示「義務化」", "SSBJ基準は2027年3月期以降段階的、現時点は義務ではない"],
        ["サプライチェーン人権DD「義務」", "経産省ガイドラインは任意"],
        ["食料・農業「2025年改正農業基本法」", "正式名称は「食料・農業・農村基本法」（2024年6月施行）"],
        ["種苗法「2025年改正完全施行」", "2020年改正種苗法は2022年4月に完全施行済み"],
        ["HACCP「2024年以降義務化」", "2021年6月に完全義務化済み"],
        ["外食へのアレルゲン義務拡大", "食品表示法は容器包装食品が対象、外食は任意"],
        ["高校金融教育「2025年義務化」", "2022年4月から高校家庭科で必修化（既に施行済み）"],
    ],
    [7.0, 11.5]
)
note("対策：法的根拠（法律名・条文・施行日）を必ず一次ソース確認。施行年がない「義務化」は原則WebSearchで確認。")

doc.add_paragraph()
heading3("δ. 法令成立日と施行日（適用日）のタイムラグ誤記")
bullet("暗号資産分離課税：法案成立は2026年3月31日、実際の適用は2028年1月1日から")
bullet("取適法（改正下請法）：2026年1月施行")
note("対策：税制改正は「成立日」「公布日」「施行日」「適用日（納税義務発生日）」を個別確認する。")

doc.add_paragraph()
heading3("ε. 義務化の段階的スケジュール誤記")
bullet("SSBJ（サステナビリティ基準）：プライム上場企業は2027年3月期から、全上場企業義務化は2031年3月期から")
bullet("人的資本開示義務：「金融商品取引法」→「企業内容等の開示に関する内閣府令」（2023年1月31日施行）が正確")
note("対策：義務化は「対象企業フェーズ×施行年」のマトリクスで確認する。")

doc.add_paragraph()
heading3("σ. 法令施行日・経過措置終了日の混同")
bullet("くるみ特定原材料義務化：施行日（2023年3月9日）と経過措置終了日（2025年3月31日）を混同し「2025年4月義務化」と誤記")
note("対策：「法律の施行日」「経過措置終了日」「完全義務化日」の3点を必ず区別して記載。")

doc.add_paragraph()
heading3("H. 法令・制度名の誤り")
add_table(
    ["誤記", "正確な名称"],
    [
        ["「医療機器リユース制度」", "「単回使用医療機器再製造制度」"],
        ["道路運送法「78条3号」", "「78条2号」（自家用有償旅客運送）が正しい根拠条文"],
        ["「J-Innovation Platform」", "「J-Innovation HUB」（経産省公式名称）"],
    ],
    [7.0, 11.5]
)
note("対策：制度名は省庁公式サイトの一次情報のみ参照する。")

# =====================================================================
heading2("【カテゴリ3】統計・数値の誤り")

heading3("C. 出典不明の具体的数値")
add_table(
    ["誤記", "正確な情報"],
    [
        ["食材ロス率「15〜30%」", "農水省調査「3〜5%」（10倍近い過大計上）"],
        ["公立小中高教員「約90万人」", "令和5年度文科省調べ「約66万人」（35%過大）"],
        ["「93〜95%のAI肌解析精度」", "根拠なし（削除）"],
        ["大学生「約300万人」", "約264万6千人（文科省2025年度・過去最多）"],
        ["糖尿病有病者・予備群「数千万人」", "約2,000万人（令和5年国民健康・栄養調査）"],
        ["介護人材不足「38万人超」", "約32万人（第8期介護保険事業計画推計）"],
    ],
    [6.5, 12.0]
)
note("対策：数値は必ず「出典名・調査年月」をセットで記載。なければ削除する。")

doc.add_paragraph()
heading3("D. 時系列の誤り・古いデータの混入")
bullet("ケアFit：「介護離職率12.4%は全産業平均14.2%を下回る」→ 14.2%は2022年の古いデータ。2024年の全産業平均は約11.5%で介護が上回っていた（方向が逆転）")
bullet("NISA口座数：「2,559万」→ 金融庁公式「2,696万口座（2025年6月末）」")
bullet("不登校：「過去最多（数値なし）」→ 35万3,970人（12年連続増・2024年度）")
note("対策：統計は調査年・公表日・改定有無を必ず確認する。")

doc.add_paragraph()
heading3("α'. 数値の「倍率」と「増加率」の混同")
bullet("「前年度比131%増」と「前年度比131%（=1.31倍）」は全く異なる")
bullet("国税庁発表の「対前年比131.4%」は「1.314倍」の意味（増加率は31.4%）")
note("対策：「前年度比△△%増」と「前年度比△△%」は区別して記述。出典の原文表現を確認する。")

doc.add_paragraph()
heading3("β'. 市場予測値の時点・出所の誤記")
bullet("睡眠市場「2030年2,162億円」→ 実際の矢野経済研究所予測は「2027年160億円」")
bullet("現在値と将来予測値、国内と国外の数値を混在させるパターンが多い")
note("対策：市場規模数値は「○○年・○○調査・国内/世界」の3軸を必ず明記する。")

doc.add_paragraph()
heading3("X. 同一アイデア内の数値矛盾（フィールド間整合性）")
bullet("whyNow「年間数万件」vs scoreComments「数千件規模」という内部矛盾")
bullet("scoreCommentsで「10倍の学習効率向上（根拠なし）」→ whyNowで「効果測定困難」と矛盾")
note("対策：同一アイデア内で同じ指標が複数フィールドに出る場合は整合性チェックを行う。")

# =====================================================================
heading2("【カテゴリ4】ハードウェア・製品情報の誤り")

heading3("T・η. ハードウェア製造終了・EOLの見落とし")
add_table(
    ["製品名", "誤記内容", "正確な情報"],
    [
        ["HoloLens 2", "「価格下落で普及」と記載", "2024年10月製造終了・2027年末サポート終了"],
        ["Meta Quest Pro", "競合リストに記載", "2024年3月販売終了（EOL）"],
        ["SoftBank Whiz", "「廃番」の記述", "2025年現在も現役販売中（かつRaaS提供中）"],
    ],
    [4.0, 5.5, 9.0]
)
note("対策：XR・ロボティクス・ウェアラブル系のハードウェアは「EOL」「discontinued」でWebSearch検索する。")

doc.add_paragraph()
heading3("V・θ. 企業・サービスのカテゴリ誤認")
bullet("Shelf Engine：「在庫管理ロボット」として記載 → 実際はAI需要予測SaaS（物理ロボットは扱わない）")
bullet("Matterport：独立企業扱い → 2025年2月にCoStar Groupに買収完了")
bullet("RT.WORKS：独立系スタートアップ → スズキ株式会社の子会社（2025年9月完全子会社化）")
note("対策：競合企業の親会社・買収情報はCrunchbase・日本経済新聞等で確認する。")

# =====================================================================
heading2("【カテゴリ5】市場・規制の新情報の見落とし")

heading3("τ. 国家・政府機関の無償OSSリリースによる競合強度急上昇")
bullet("デジタル庁ガバメントAI「源内」：2026年4月24日にMITライセンスでOSS公開")
bullet("政府OSSは「価格競合ゼロ」で市場に参入するため差別化軸の再定義が必要")
note("対策：デジタル庁・IPAのOSS公開情報を定期的に確認する。")

doc.add_paragraph()
heading3("φ. 業界大手のAI機能急拡充による差別化軸侵食")
bullet("Mazrica Target（2025年11月）：550万社DBから個社特化提案書1クリック生成")
bullet("スタディング（2024-2025年）：AI実力スコア・AI学習ナビ・AI学習プランオプティマを大幅強化")
bullet("Jasper（2025年）：Workflows機能でSEO戦略立案まで対応")
note("対策：競合の機能更新は公式ブログ・PR TIMESで「AI 新機能 2025」として定期監査（6ヶ月ごと）。")

doc.add_paragraph()
heading3("ν. SaaS競合の事業売却・大手傘下化の未反映")
add_table(
    ["サービス", "変更内容", "時期"],
    [
        ["Udemy", "Courseraとの合併発表", "2025年12月（2026年下半期完了予定）"],
        ["資格スクエア", "学研ホールディングス傘下（レアジョブから事業売却）", "2026年1月"],
        ["Classi", "ベネッセコーポレーションに吸収合併", "2026年4月1日"],
        ["スタディサプリ", "学習塾向けサービス終了（個人向けは継続）", "2026年3月30日"],
    ],
    [4.5, 10.0, 4.0]
)
note("対策：EdTech・SaaS競合は「事業売却」「合併」「サービス終了」を定期的にWebSearchで確認。")

# =====================================================================
heading2("【カテゴリ6】インサイトパネル専用チェック（whyNow・noveltyNote）")

heading3("ω. 廃止サービスがnoveltyNote内に競合として残存")
bullet("ペットみるん（東京電力EP）：2020年5月31日サービス終了済みなのに競合として記載")
bullet("noveltyNoteの「〇〇は存在しない」「〇〇市場は空白」という記述を検証する際、その根拠として使っているサービスも存在確認必須")

doc.add_paragraph()
heading3("γ'. 法令の根拠条文・命令名の誤記（インサイトパネル内）")
bullet("人的資本開示義務の根拠：「金融商品取引法」ではなく「企業内容等の開示に関する内閣府令」")
bullet("法律の改正と省令・内閣府令の改正を混同するパターン")
note("対策：法令記述は「法律名」「省令/政令/内閣府令名」「施行日」の3点セットで確認。")

doc.add_paragraph()
heading3("strengthNote専用チェックの知見")
body("全200件のstrengthNoteに専用チェックを実施した結果：", bold=True)
bullet("「データフライホイール」「スイッチングコスト」「ネットワーク効果」「先行者優位」等の抽象的優位性記述が中心")
bullet("抽象的優位性の主張はファクトチェック可能な具体的事実を含まないため、誤り率がほぼゼロ")
bullet("strengthNoteは今後のファクトチェックの優先度を下げてよい（抽象記述なので誤り率が構造的に低い）")
bullet("限られた監査リソースはwhyNow・noveltyNote・competitors・scoreCommentsに集中すべき")

# =====================================================================
pagebreak()
heading1("4. 監査の教訓・改善提言")

body("以下の7つの提言を実装することで、初期生成時の品質を大幅に向上できる。", bold=True)
doc.add_paragraph()

add_table(
    ["番号", "提言", "詳細"],
    [
        ["1", "アイデア生成時にWebSearch必須化",
         "生成段階でリアルタイム検索を組み込めば初期品質が大幅向上。最重要施策。"],
        ["2", "competitors は生成後即時確認",
         "「廃止済みサービスを競合に書く」は最多ミス。サービス稼働確認の自動化が有効。"],
        ["3", "数値には必ず出典注記を",
         "生成AIは出典なしに数値を混入させやすい。「数値を書くなら出典を書く」をルール化。"],
        ["4", "法規制フィールドは更新トリガーを設ける",
         "義務化・施行日・適用範囲は変わる。法改正ニュースをトリガーに自動フラグ立てを検討。"],
        ["5", "competitors フィールドは6ヶ月ごとに再生成推奨",
         "競合状況は最も陳腐化が早い。定期自動更新の仕組みが必要。"],
        ["6", "「義務化」「唯一」「急増」はハルシネーションのシグナル",
         "これらのキーワードが含まれるフィールドを優先的にファクトチェック。"],
        ["7", "国産SaaS競合は日本語検索で追跡",
         "英語圏の情報に出にくい。日本語での「競合サービス名 + 2024 or 2025」検索を必ず実施。"],
    ],
    [1.0, 4.5, 13.0]
)

# =====================================================================
heading1("5. 収穫逓減モデルと監査計画")

heading2("5-1. ラウンド別修正件数の推移")
body("繰り返しチェックには収穫逓減の傾向がある。監査リソースを最適配分する参考指標：")
add_table(
    ["ラウンド", "修正件数の目安", "主な発見内容"],
    [
        ["R1〜R3", "1ラウンドで10〜20件", "廃止サービス・義務化誤用・統計数値の桁違い等の基本的事実誤認"],
        ["R4〜R6", "1ラウンドで5〜10件", "リブランド未反映・機能拡張の見落とし・細部の精度向上"],
        ["R7〜R8", "1ラウンドで3〜7件", "新視点（価格変化・AI機能拡充・統計更新）による追加誤り発見"],
        ["R9以降", "1ラウンドで0〜4件", "マイナーな数値精度・出典記載の改善（限界収益逓減）"],
    ],
    [2.5, 4.5, 11.5]
)

heading2("5-2. 推奨監査スケジュール")
bullet("初回生成後：R1〜R3を1日で集中実施（最大効果）")
bullet("3ヶ月後：competitors・whyNow・noveltyNoteの重点再監査")
bullet("6ヶ月後：全フィールド再監査（AIサービス価格変動・法令改正を中心に）")
bullet("年次：市場統計全数値の年次更新（調査年を「○○年調査」と明記）")

heading2("5-3. 各ラウンドで設定すべき「新視点」")
body("同じ観点で繰り返すと見落としが増える。毎ラウンド以下の視点を追加・変更する：", bold=True)
bullet("R4：競合AI機能の急拡張・法令施行タイムラグ・ハードウェアEOL")
bullet("R5：フィールド間整合性・架空サービス名・地域展開・機能限定の誤誇張")
bullet("R6：数値単位ミス・企業経営状況変化・trendKeywords陳腐化")
bullet("R7：競合サービスのリブランド・廃止・合併・子会社化")
bullet("R8：サービス価格変化・大企業新規参入・AI機能急拡充・統計数値鮮度")

# =====================================================================
pagebreak()
heading1("6. 確認済み重要事実一覧（2026年4月時点）")

body("以下の事実は過去の監査で確認・修正済み。新規アイデア生成時の参照データとして活用する。")
doc.add_paragraph()

heading2("6-1. サービス・企業情報")
add_table(
    ["事実", "詳細"],
    [
        ["MATTERPORT", "2025年2月28日にCoStar Groupに約16億ドルで買収完了"],
        ["GVA assist", "2024年11月に「OLGA」へ統合リブランド"],
        ["カレコ", "2024年2月20日に「三井のカーシェアーズ」へ正式改称"],
        ["FiNC", "2026年9月末サービス終了予定（2026年4月1日より新規受付停止）"],
        ["NTTコムオンライン", "2026年1月1日付で「NTTドコモビジネスX株式会社」に社名変更"],
        ["PayPay資産運用", "2026年2月にミニアプリ名称「PayPay証券」へ変更"],
        ["ウェルスナビ", "MUFGのeスマート証券（旧auカブコム証券）と合併発表（2027年度中予定・「エムット」ブランド）"],
        ["Udemy-Coursera合併", "2026年4月9日に株主99%超が承認・Q2クローズ目標"],
        ["Spotify for Creators", "2024年11月13日にSpotify for Podcastersからリブランド"],
        ["Zendesk", "2026年3月にForethought（自律型AIカスタマーサービス）を買収"],
        ["デジタル庁ガバメントAI「源内」", "2026年4月24日OSSリリース（MITライセンス）"],
        ["Mazrica Target", "2025年11月1日リリース、550万社企業DB＋Web情報→提案書1クリック生成"],
        ["Genspark", "2026年1月に日本法人設立・Plusプラン（月24.99ドル）"],
    ],
    [5.5, 13.0]
)

doc.add_paragraph()
heading2("6-2. 法令・制度")
add_table(
    ["事実", "詳細"],
    [
        ["暗号資産分離課税", "法案成立2026年3月31日、適用は2028年1月1日から"],
        ["食品接触材ポジティブリスト制度", "2025年6月完全施行"],
        ["カシューナッツ義務化", "2026年4月1日から特定原材料義務化（9品目体制へ移行）"],
        ["取適法（改正下請法）", "2026年1月施行"],
        ["SSBJ義務化", "プライム市場2027年3月期→全上場2031年3月期（段階的）"],
        ["農地バンク経由一本化", "2025年4月の農業経営基盤強化促進法改正で農地貸借が一本化"],
        ["フリーランス保護新法", "2024年11月1日施行"],
        ["育児・介護休業法改正", "2025年4月より企業に介護前情報提供・個別周知・意向確認が義務化"],
        ["副業・兼業促進ガイドライン最新改訂", "2025年3月31日版"],
    ],
    [5.5, 13.0]
)

doc.add_paragraph()
heading2("6-3. 統計数値")
add_table(
    ["統計項目", "確認済み最新値", "出典"],
    [
        ["日銀政策金利", "0.75%（2025年12月・1995年以来30年ぶり高水準）", "日本銀行"],
        ["NISA口座数", "2,696万口座（2025年6月末）", "金融庁"],
        ["訪日外客数", "4,268万人（2025年・過去最高）", "JNTO確定値"],
        ["不登校", "35万3,970人（小中学生・12年連続増）", "文科省令和6年度調査"],
        ["大学生数", "264万6千人（2025年度・過去最多）", "文科省"],
        ["最低賃金（2025年度）", "全国加重平均1,121円・引き上げ幅66円", "厚生労働省確定値"],
        ["建設業就業者数", "2024年477万人（1997年ピーク685万人から激減）", "建設業統計"],
        ["小学校数", "18,822校（国公私立合計・2024年）", "文科省学校基本調査"],
        ["2025年M&A件数", "5,115件（過去最多）", "レコフデータ"],
        ["国内アパレル市場", "8兆5,010億円（2024年実績）", "矢野経済研究所"],
        ["SleepioRx承認種別", "FDA「クリアランス（510K）」（「認可（Approval）」ではない）", "FDA公式"],
    ],
    [5.0, 7.0, 6.5]
)

# =====================================================================
# フッター
# =====================================================================
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("本ドキュメントは AideaSpark（paint-platform.com/aideaspark）の監査活動に基づき作成。最終更新：2026年4月28日")
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# 保存
doc.save(OUTPUT_FILE)
print("OK: " + OUTPUT_FILE)
