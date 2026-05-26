"""
AideaSpark アイデア生成パイプライン ドキュメント生成スクリプト
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---- スタイル設定 ----
style_normal = doc.styles['Normal']
style_normal.font.name = 'Meiryo UI'
style_normal.font.size = Pt(10.5)

def set_font(run, bold=False, size=None, color=None):
    run.font.name = 'Meiryo UI'
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Meiryo UI'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x27, 0x27, 0x27)
        elif level == 3:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent * 0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    return p

def add_bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p

def add_note(text, kind='info'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    prefix = '⚠ ' if kind == 'warn' else '✅ ' if kind == 'ok' else '※ '
    color = (0xd9, 0x53, 0x00) if kind == 'warn' else (0x0a, 0x74, 0x40) if kind == 'ok' else (0x44, 0x44, 0x66)
    run = p.add_run(prefix + text)
    set_font(run, bold=(kind == 'warn'), color=color)
    return p

# ========== 表紙 ==========
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('AideaSpark')
run.font.name = 'Meiryo UI'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('アイデア生成パイプライン ドキュメント')
run.font.name = 'Meiryo UI'
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'作成日: {datetime.date.today().strftime("%Y年%m月%d日")}  /  対象プロジェクト: C:\\Users\\after\\bizidea')
run.font.name = 'Meiryo UI'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_page_break()

# ========== 1. 概要 ==========
add_heading('1. 概要', 1)
add_body('AideaSpark は Next.js 16 + Supabase/PostgreSQL + Claude API で構成されるビジネスアイデア探索プラットフォームである。'
         'アイデアコンテンツは「バッチ生成（スクリプト）」と「リアルタイム生成（API）」の2経路で生成・管理されている。'
         '現在のアイデア総数は 200 件（idea-001〜idea-100 + idea-101〜idea-200）。')

add_heading('1.1 全体データフロー', 2)
flow_lines = [
    ('Step 1', 'batch-generate-ideas.mjs', 'Claude API + Web Search → scripts/generated/*.json に保存'),
    ('Step 2', 'import-generated.mjs',     'JSON → src/data/mock/ideas.ts に TypeScript 形式で追記'),
    ('Step 3', 'generate-insights.mjs',    'ideas.ts をパース → whyNow / noveltyNote / strengthNote を生成・書き戻し'),
    ('Step 4', 'prisma/seed.ts',           'ideas.ts → Supabase (PostgreSQL) へ upsert'),
    ('Step 5', 'apply-insights.mjs',       'insights-data*.json → Supabase へ直接 UPDATE (別経路で生成した場合)'),
    ('Step 6', 'assign-patterns.cjs',      '既存20件のパターンを PostgreSQL へ直接 UPDATE (補完用)'),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for cell, txt in zip(hdr, ['ステップ', 'スクリプト名', '処理内容']):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Meiryo UI'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for step, script, desc in flow_lines:
    row = tbl.add_row().cells
    for cell, txt in zip(row, [step, script, desc]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Meiryo UI'
            run.font.size = Pt(9.5)

doc.add_paragraph()

# ========== 2. スクリプト詳細 ==========
add_heading('2. スクリプト詳細', 1)

# 2.1
add_heading('2.1  batch-generate-ideas.mjs　—　メイン生成エンジン', 2)
add_body('アイデアを大量にカテゴリ別・バッチ単位で生成するコアスクリプト。')
add_bullet('使用モデル: claude-sonnet-4-6')
add_bullet('処理フロー:')
add_bullet('① カテゴリ別に既存サービス名を収集（重複防止）', level=2)
add_bullet('② Web Search でカテゴリ関連の規制・競合・市場規模を取得', level=2)
add_bullet('③ 市場シグナル付きプロンプトで 10 件のアイデアを生成', level=2)
add_bullet('④ JSON 抽出 → ID・番号の連番保証 → scripts/generated/ に保存', level=2)
add_bullet('出力: scripts/generated/ph1_{カテゴリ}.json（例: ph1_SaaS.json）')
add_bullet('パターン: 74 パターン（A〜J の 10 カテゴリ）がプロンプトに組み込まれ、各アイデアに 2〜3 個を付与')
add_bullet('生成制約:')
add_bullet('架空企業・一般化禁止（実在企業のみ記載）', level=2)
add_bullet('数値は「○○%」等の定量表現を原則使用しない', level=2)
add_bullet('スコア全項目は異なる値（max-min ≥ 1）', level=2)
add_bullet('各フィールドに指示文・説明文の混入禁止', level=2)

add_heading('プロンプト構造（概略）', 3)
add_code('Phase 1: Web Search\n'
         '  "日本市場の「{カテゴリ}」領域について：\n'
         '   1. 主要プレイヤー・競合サービス\n'
         '   2. 市場規模・成長率\n'
         '   3. 規制変化・政策動向\n'
         '   4. 未解決課題"\n\n'
         'Phase 2: アイデア生成\n'
         '  system: SYSTEM_PROMPT（74パターン定義 + 生成制約）\n'
         '  user:   buildPrompt(category, count, existingNames, startId, searchContext)')

doc.add_paragraph()

# 2.2
add_heading('2.2  import-generated.mjs　—　JSON → TypeScript 変換', 2)
add_body('生成済み JSON を src/data/mock/ideas.ts に追記するスクリプト。')
add_bullet('入力: scripts/generated/*.json')
add_bullet('処理: 重複 ID チェック → TypeScript オブジェクト文字列に変換 → ideas.ts 末尾の ]; 直前に挿入')
add_bullet('出力: src/data/mock/ideas.ts（追記）')
add_bullet('次ステップ: npx tsx prisma/seed.ts を手動実行')

doc.add_paragraph()

# 2.3
add_heading('2.3  generate-insights.mjs　—　インサイト一括生成', 2)
add_body('全アイデアに 3 つの深掘りインサイトフィールドを付与するスクリプト。')
add_bullet('対象フィールド: whyNow / noveltyNote / strengthNote')
add_bullet('バッチサイズ: 5 件 / API 呼び出し（レート制限対応）')
add_bullet('進捗管理: insights-progress.json で中断後の再開に対応')

add_heading('生成ルール', 3)
rules = [
    ('whyNow', '2文。規制変化・技術転換点・行動変容のうち最低1つを引用し、具体的な数値・年月を含める'),
    ('noveltyNote', '2文。「競合○○は〜だが△△という点が存在しない」対比構造。競合名の明記必須'),
    ('strengthNote', '2文。収益構造・参入障壁・ネットワーク効果・スイッチングコストを具体的に記述'),
]
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
for cell, txt in zip(tbl2.rows[0].cells, ['フィールド', '生成ルール']):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Meiryo UI'
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for field, rule in rules:
    row = tbl2.add_row().cells
    for cell, txt in zip(row, [field, rule]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Meiryo UI'
            run.font.size = Pt(9.5)
doc.add_paragraph()

add_heading('2.4  prisma/seed.ts　—　DB 投入', 2)
add_body('ideas.ts をパースして Supabase (PostgreSQL) に upsert するメインシードスクリプト。')
add_bullet('接続先: DATABASE_URL 環境変数で指定（.env）')
add_bullet('処理: prisma.idea.upsert — id をキーに create / update')
add_bullet('投入フィールド: serviceName, concept, target, problem, product, revenueModel, competitors, competitiveEdge, '
           'tags, category, targetIndustry, targetCustomer, investmentScale, difficulty, scores, scoreComments, '
           'trendKeywords, patterns, whyNow, noveltyNote, strengthNote, patternRationale 他')
add_note('seed.ts は idea の全フィールドを上書きする。事後に apply-insights.mjs や assign-patterns.cjs で'
         '個別フィールドを UPDATE する場合は seed.ts の再実行で上書きされる点に注意。', kind='warn')

doc.add_paragraph()

add_heading('2.5  apply-insights.mjs　—　インサイト別経路 DB 反映', 2)
add_body('scripts/insights-data*.json（事前生成済みインサイト）を pg 直接接続で Supabase に UPDATE するスクリプト。')
add_bullet('入力: scripts/insights-data.json / insights-data-2.json / insights-data-3.json（計 100 件: idea-001〜100）')
add_bullet('処理: UPDATE "Idea" SET "whyNow"=$1, "noveltyNote"=$2, "strengthNote"=$3 WHERE id=$4')
add_bullet('用途: generate-insights.mjs を経由せずに事前生成データを直接適用する場合')
add_note('2026-04-28 に初めて実行し、idea-001〜100 のインサイトを Supabase に反映済み。', kind='ok')

doc.add_paragraph()

add_heading('2.6  assign-patterns.cjs　—　パターン補完', 2)
add_body('初期 20 件（idea-001〜020）のパターン ID をハードコードで直接 UPDATE するスクリプト。')
add_bullet('方式: ハードコードされたマッピング（20 アイデア × 2〜3 パターン）')
add_bullet('例: idea-001 (TASQ) → ["B-1", "F-1"]')
add_bullet('接続: pg.Pool + DATABASE_URL で PostgreSQL に直接 UPDATE')

doc.add_paragraph()

# ========== 3. アイデアデータの現状 ==========
add_heading('3. アイデアデータの現状', 1)

add_heading('3.1 件数・ファイル構成', 2)
add_bullet('src/data/mock/ideas.ts: 200 件（6,002 行）— 単一 TypeScript 配列として管理')
add_bullet('scripts/generated/: バッチ生成 JSON（ph1_SaaS.json / ph1_D2C.json 等 8 ファイル）')
add_bullet('scripts/insights-data*.json: インサイト別保管 JSON（idea-001〜100、計 100 件）')

add_heading('3.2 フィールド充足状況', 2)

tbl3 = doc.add_table(rows=1, cols=6)
tbl3.style = 'Table Grid'
headers = ['世代', 'ID 範囲', 'patterns', 'patternRationale', 'whyNow / noveltyNote\n/ strengthNote', '備考']
for cell, txt in zip(tbl3.rows[0].cells, headers):
    cell.text = txt
    p = cell.paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.name = 'Meiryo UI'
    p.runs[0].font.size = Pt(9)

rows_data = [
    ('Phase 1 (手動)', 'idea-001〜020', '✅', '✅', '✅ (DB反映済)', '初期ハードコード分'),
    ('Phase 2 (バッチ)', 'idea-021〜100', '✅', '❌ 欠落', '✅ (2026-04-28 DB反映)', 'インサイトは apply-insights.mjs で反映'),
    ('Phase 2 (バッチ)', 'idea-101〜200', '✅', '❌ 欠落', '✅ (ideas.ts 埋め込み済)', 'seed.ts で DB に入っている'),
]
for cols in rows_data:
    row = tbl3.add_row().cells
    for cell, txt in zip(row, cols):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Meiryo UI'
            run.font.size = Pt(9)

doc.add_paragraph()

# ========== 4. リアルタイム生成（API） ==========
add_heading('4. リアルタイム生成 API（/api/ai-generate）', 1)
add_body('ユーザーが UI 上でアイデア生成を依頼した場合に動作する 3 段階処理。バッチ生成とは別体系。')

phases = [
    ('Phase 1', 'Web Search', 'ユーザー指定テーマで市場調査（競合・規制・市場規模）'),
    ('Phase 2', 'アイデア生成', '検索結果 + 市場シグナル付きプロンプトでアイデア生成'),
    ('Phase 3', 'ファクトチェック', 'competitors 検証・具体数値確認・矛盾修正'),
]
for phase, name, desc in phases:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{phase}: {name}　')
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_note('ユーザー生成アイデアは DB に保存されない（一時的なレスポンスのみ）。GenerationLog には記録される。', kind='warn')
add_note('バッチ生成分（idea-001〜200）にはこの 3 段階ファクトチェックが適用されていない。', kind='warn')

doc.add_paragraph()

# ========== 5. パターン体系 ==========
add_heading('5. パターン体系', 1)
add_body('74 パターン（A〜J の 10 カテゴリ）によりアイデアを戦略的に分類。各アイデアに 2〜3 パターンを付与。')

patterns_table = [
    ('A', '顧客・市場深掘り型', 'JTBD、ノンカスタマー分析、セグメント再定義 など'),
    ('B', 'テクノロジー先行型', '先端技術転用、AI×XR×IoT 融合、バーティカル AI など'),
    ('C', 'ビジネスモデル革新型', 'サブスク、プレミアム、マーケットプレイス、API 化 など'),
    ('D', 'バリューチェーン変革型', '中抜き、川上統合、川下拡張 など'),
    ('E', 'エコシステム型', 'プラットフォーム、ネットワーク効果、コミュニティ など'),
    ('F', '規制・社会変化対応型', '規制変化の窓、社会課題解決、行動変容の波 など'),
    ('G', 'グローバル・クロスボーダー型', '海外展開、逆輸入、ローカル×グローバル など'),
    ('H', 'データ活用型', 'データ収益化、フライホイール、AIループ など'),
    ('I', 'コスト構造革新型', 'CAPEX→OPEX、ゼロ限界費用、アウトソーシング など'),
    ('J', 'ブランド・コミュニティ型', 'コミュニティ主導、ファン経済、ストーリー販売 など'),
]
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
for cell, txt in zip(tbl4.rows[0].cells, ['カテゴリ', '名称', '代表パターン例']):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Meiryo UI'
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for cat, name, examples in patterns_table:
    row = tbl4.add_row().cells
    for cell, txt in zip(row, [cat, name, examples]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Meiryo UI'
            run.font.size = Pt(9.5)
doc.add_paragraph()

# ========== 6. 既知の課題 ==========
add_heading('6. 既知の課題', 1)

issues = [
    ('❌', 'patternRationale の欠落',
     'idea-021〜200（バッチ生成分）は patternRationale（パターン選択の根拠説明）が未入力。'
     'Prisma スキーマにはカラムが存在するが、generate-insights.mjs の生成対象に含まれていない。'),
    ('❌', 'バッチ生成済みデータのファクトチェック未実施',
     '/api/ai-generate の 3 段階ファクトチェックはリアルタイム生成のみに適用される。'
     'idea-001〜200 のバッチ生成分には事後ファクトチェックが行われておらず、'
     'competitors の廃止サービスや古い規制情報が混入している可能性がある。'
     '（R1〜R5 として idea-101〜200 の手動ファクトチェックを一部実施済み）'),
    ('❌', 'パターン選択の再現性なし',
     'バッチ生成時のパターン付与は LLM の任意判断によるため、同一テーマを再実行しても'
     '異なるパターンが選ばれる可能性がある。検証・固定化の仕組みがない。'),
    ('❌', 'ユーザー生成アイデアの未保存',
     '/api/ai-generate で生成されたアイデアは DB に保存されない。'
     'GenerationLog には記録されるが、フィードには反映されないため学習機会が失われている。'),
    ('❌', 'ID 管理の手動化',
     '新カテゴリ追加時は startId を手動で計算する必要があり、連番衝突のリスクがある。'
     '自動連番管理の仕組みが存在しない。'),
    ('❌', 'インサイトとパターンの整合性検証なし',
     'whyNow で「規制変化」を述べても patterns に F 系（規制対応型）が含まれているか未検証。'
     'コンテンツの整合性チェック機構が存在しない。'),
    ('⚠', 'seed.cjs（SQLite 版）と seed.ts（Prisma 版）の二重管理',
     'prisma/seed.cjs は dev.db（ローカル SQLite）向け、prisma/seed.ts は Supabase 向け。'
     '前者は whyNow/noveltyNote/strengthNote/patterns フィールドを INSERT していない（列定義が古い）。'),
]

for icon, title, desc in issues:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{icon}  {title}\n')
    set_font(r1, bold=True, color=(0xc0, 0x39, 0x2b) if icon == '❌' else (0xe6, 0x7e, 0x22))
    r2 = p.add_run(desc)
    set_font(r2, color=(0x33, 0x33, 0x33))
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# ========== 7. 環境・接続情報 ==========
add_heading('7. 環境・接続情報', 1)
add_bullet('プロジェクトルート: C:\\Users\\after\\bizidea')
add_bullet('フレームワーク: Next.js 16 + TypeScript + shadcn/ui')
add_bullet('DB: Supabase (PostgreSQL) — Prisma ORM 経由')
add_bullet('ローカル SQLite: prisma/dev.db（seed.cjs 専用。インサイト列なし）')
add_bullet('Claude API: claude-sonnet-4-6（バッチ生成・インサイト生成・リアルタイム生成共通）')
add_bullet('開発サーバーポート: 4001')
add_bullet('主要スクリプト実行方法:')
add_code('# バッチ生成\nnode scripts/batch-generate-ideas.mjs --phase ph1 --category SaaS --count 10\n\n'
         '# TypeScript 化\nnode scripts/import-generated.mjs\n\n'
         '# インサイト生成\nnode scripts/generate-insights.mjs\n\n'
         '# DB 投入（Supabase）\nnpx tsx prisma/seed.ts\n\n'
         '# インサイト別経路反映\nnode scripts/apply-insights.mjs\n\n'
         '# 監査（品質チェック）\nnode scripts/audit-insights.mjs')

doc.add_paragraph()

# ========== フッター ==========
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(f'本ドキュメントは {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")} 時点のコードベース調査に基づき自動生成されました。')
set_font(run, color=(0x99, 0x99, 0x99))
run.font.size = Pt(9)

# ========== 保存 ==========
output_path = r'C:\Users\after\AI副業_YouTube\AideaSpark_パイプライン概要.docx'
doc.save(output_path)
import sys
sys.stdout.buffer.write(f'[OK] 出力完了: {output_path}\n'.encode('utf-8'))
